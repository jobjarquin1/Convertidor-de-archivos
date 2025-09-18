#!/usr/bin/env python3
# file_converter_gui.py


import os
from tkinter import Tk, Label, Button, filedialog, messagebox, StringVar, OptionMenu
from pathlib import Path

# Librerías de conversión
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Preformatted, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from pdfminer.high_level import extract_text
import markdown

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except Exception:
    PANDOC_AVAILABLE = False


def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def write_txt(path, text):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)

def read_docx(path):
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def write_docx(path, text):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(path)

def read_pdf_text(path):
    return extract_text(path)

def write_pdf_from_text(path, text):
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Preformatted(text, styles['Code']), Spacer(1,12)]
    doc.build(story)

def md_to_txt(src, dst):
    md = read_txt(src)
    html = markdown.markdown(md)
    from html.parser import HTMLParser
    class TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.result = []
        def handle_data(self, data):
            self.result.append(data)
    parser = TextExtractor()
    parser.feed(html)
    text = "\n".join(parser.result)
    write_txt(dst, text)

def convert_file(src, dst, target_format):
    ext = Path(src).suffix.lower().lstrip(".")
    if ext == "docx":
        if target_format=="pdf":
            write_pdf_from_text(dst, read_docx(src))
        elif target_format=="txt":
            write_txt(dst, read_docx(src))
    elif ext=="txt":
        if target_format=="docx":
            write_docx(dst, read_txt(src))
        elif target_format=="pdf":
            write_pdf_from_text(dst, read_txt(src))
    elif ext=="pdf":
        if target_format=="txt":
            write_txt(dst, read_pdf_text(src))
    elif ext=="md":
        if target_format=="txt":
            md_to_txt(src, dst)
        elif target_format=="docx":
            if PANDOC_AVAILABLE:
                pypandoc.convert_file(src, "docx", outputfile=dst)
            else:
                md_to_txt(src, dst + ".tmp.txt")
                write_docx(dst, read_txt(dst + ".tmp.txt"))
                os.remove(dst + ".tmp.txt")
        elif target_format=="pdf":
            if PANDOC_AVAILABLE:
                pypandoc.convert_file(src, "pdf", outputfile=dst)
            else:
                md_to_txt(src, dst + ".tmp.txt")
                write_pdf_from_text(dst, read_txt(dst + ".tmp.txt"))
                os.remove(dst + ".tmp.txt")
    else:
        messagebox.showerror("Error", f"Conversión desde {ext} no soportada")
        return
    messagebox.showinfo("Listo", f"✅ Archivo convertido: {dst}")


class FileConverterGUI:
    def __init__(self, master):
        self.master = master
        master.title("Convertidor de Archivos (DOCX/PDF/TXT/MD)")

        Label(master, text="Selecciona archivo:").grid(row=0, column=0, padx=10, pady=10)
        Button(master, text="Abrir archivo", command=self.browse_file).grid(row=0,column=1,padx=10)

        Label(master, text="Formato destino:").grid(row=1,column=0,padx=10,pady=10)
        self.target_format = StringVar(master)
        self.target_format.set("pdf")
        OptionMenu(master, self.target_format, "pdf","txt","docx").grid(row=1,column=1,padx=10)

        Button(master, text="Convertir", command=self.convert_action).grid(row=2,column=0,columnspan=2,pady=20)

        self.src_path = None

    def browse_file(self):
        file_path = filedialog.askopenfilename(title="Seleccionar archivo", filetypes=[("Todos los archivos","*.*")])
        if file_path:
            self.src_path = file_path
            messagebox.showinfo("Archivo seleccionado", file_path)

    def convert_action(self):
        if not self.src_path:
            messagebox.showerror("Error", "Primero selecciona un archivo")
            return
        src = self.src_path
        tgt = self.target_format.get()
        dst = str(Path(src).with_suffix(f".{tgt}"))
        convert_file(src,dst,tgt)

if __name__ == "__main__":
    root = Tk()
    app = FileConverterGUI(root)
    root.mainloop()
